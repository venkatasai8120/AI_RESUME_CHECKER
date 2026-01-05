from __future__ import annotations

import os
import tempfile
from typing import Tuple

from docx import Document
from PyPDF2 import PdfReader

ALLOWED_EXTENSIONS = {".pdf", ".docx"}
MAX_FILE_MB = 5


def _ext(filename: str) -> str:
    return os.path.splitext(filename or "")[1].lower()


def validate_file(filename: str, file_size_bytes: int) -> Tuple[bool, str]:
    """
    Basic safety checks: extension + size.
    """
    ext = _ext(filename)
    if ext not in ALLOWED_EXTENSIONS:
        return False, "Only .pdf and .docx files are supported."

    if file_size_bytes > MAX_FILE_MB * 1024 * 1024:
        return False, f"File too large. Max allowed is {MAX_FILE_MB} MB."

    return True, ""


def extract_text_from_docx(path: str) -> str:
    doc = Document(path)
    parts = []

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    # Also capture tables if present
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def extract_text_from_pdf(path: str) -> str:
    reader = PdfReader(path)
    parts = []
    for page in reader.pages:
        text = (page.extract_text() or "").strip()
        if text:
            parts.append(text)
    return "\n".join(parts).strip()


def extract_resume_text(path: str) -> str:
    """
    Detect file type and extract resume text.
    """
    ext = _ext(path)
    if ext == ".pdf":
        return extract_text_from_pdf(path)
    if ext == ".docx":
        return extract_text_from_docx(path)
    raise ValueError("Unsupported file type.")


def extract_text_from_uploaded_file(file_storage) -> str:
    """
    Accepts a Werkzeug FileStorage (Flask upload).
    Saves to a temp file, validates, extracts text, then deletes temp file.
    """
    filename = file_storage.filename or ""
    # Try to get size (may be 0 on some servers). We'll also validate after saving.
    try:
        pos = file_storage.stream.tell()
        file_storage.stream.seek(0, os.SEEK_END)
        size = file_storage.stream.tell()
        file_storage.stream.seek(pos, os.SEEK_SET)
    except Exception:
        size = 0

    ok, msg = validate_file(filename, size if size else 0)
    if not ok:
        raise ValueError(msg)

    ext = _ext(filename)

    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            tmp_path = tmp.name
            file_storage.save(tmp_path)

        # Validate actual saved size too
        actual_size = os.path.getsize(tmp_path)
        ok2, msg2 = validate_file(filename, actual_size)
        if not ok2:
            raise ValueError(msg2)

        return extract_resume_text(tmp_path)

    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except Exception:
                pass
